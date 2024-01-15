import os
import shutil
from docx import Document
from text_extraction import *

from setup import SOURCE_DIR, PDFs_DIR_PATH


def copy_pdf_files(src_dir: str, dest_dir: str) -> None:
    """
    Copies all PDF files from a source directory (including its subdirectories)
    to a destination directory. It creates the destination directory if it does not already exist.
    It also avoids overwritting . 

    Parameters:
    src_dir (str): Source path where we take the pdfs.
    dest_dir (str): Destination path where we send the pdfs.
    """

    if not os.path.exists(dest_dir):
        os.makedirs(dest_dir)

    # Iterate through the source directory, including all subdirectories
    for root, dirs, files in os.walk(src_dir):
        for file_name in files:
            if file_name.endswith('.pdf'):
                source = os.path.join(root, file_name)
                destination = os.path.join(dest_dir, file_name)

                counter = 0
                while os.path.exists(destination):
                    counter += 1
                    base_name, extension = os.path.splitext(file_name)
                    destination = os.path.join(dest_dir, f'{base_name}_{counter}{extension}')

                shutil.copy2(source, destination)
                print(f'Copied {source} to {destination}')


def keep_invoices_only() -> None:
    """
    This function iterates through all files in the 'PDFs_DIR_PATH' directory, identifies PDF files that are 
    either English or Spanish invoices based on specific keywords, and counts them. Files that are not identified 
    as invoices are deleted from the directory.
    """

    files = [f for f in os.listdir(PDFs_DIR_PATH) if os.path.isfile(os.path.join(PDFs_DIR_PATH, f))]

    invoice_english_counter = 0
    invoice_spanish_counter = 0
    others_counter = 0
    counter = 0
    for f in files:
        pdf_text = get_text_from_pdf(f)
        lines_pdf = re.split('\n', pdf_text)
        # Remove all whitespaces from all strings
        lines_pdf = [''.join(l.split()) for l in lines_pdf]

        isInvoice1 = any('commercialinvoice' in l.lower() for l in lines_pdf)
        isInvoice2 = any('paymentconditions' in l.lower() for l in lines_pdf)
        isInvoice3 = any('facturacomercial' in l.lower() for l in lines_pdf)
        isInvoice4 = any('condicionesdepago' in l.lower() for l in lines_pdf)

        if isInvoice1 and isInvoice2:
            invoice_english_counter += 1
        elif isInvoice3 and isInvoice4:
            invoice_spanish_counter += 1
        else:
            others_counter += 1
            file_path = PDFs_DIR_PATH + f
            if os.path.exists(file_path):
                os.remove(file_path)
                print(f'{file_path} has been deleted!')
            else:
                print(f'The file {file_path} does not exist!')

        counter += 1
        print(f'[{counter}/{len(files)}] - file analyzed')

    print(f'Invoices English: {invoice_english_counter}')
    print(f'Invoices Spanish: {invoice_spanish_counter}')
    print(f'Others {others_counter}')


def preprocess_pdf(pdf_file_name: str) -> None:
    """
    Preprocesses a given PDF file. It looks for pages containing key phrases in English and Spanish. 
    If these phrases are found in the first page, the original PDF is kept as is. 
    If found in subsequent pages, only the page where these are found and the following pages 
    are retained in a new PDF. The processed PDF is saved with the same name in the same directory.

    Parameters:
    pdf_file_name (str): The name of the PDF file.
    """

    pdf_file_path = PDFs_DIR_PATH + pdf_file_name
    output_pdf_file_path = pdf_file_path

    try:
        with open(pdf_file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            writer = PyPDF2.PdfWriter()

            # Iterate through the pages to find the target page
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                text = page.extract_text().lower()  # Get the text content in lower case
                text = ''.join(text.split())  # Remove all whitespaces

                # Check if both "commercial invoice" and "payment conditions" are present
                is_english = 'commercialinvoice' in text and 'paymentconditions' in text
                is_spanish = 'facturacomercial' in text and 'condicionesdepago' in text
                if is_english or is_spanish:
                    # If found in the first page, just save the original PDF
                    if page_num == 0:
                        writer = reader
                        return
                    else:
                        # If found in other pages, add the current and following pages to the new PDF
                        for remaining_page_num in range(page_num, len(reader.pages)):
                            writer.add_page(reader.pages[remaining_page_num])
                    break

            with open(output_pdf_file_path, 'wb') as output_file:
                writer.write(output_file)

            print(f"Processed file saved to {output_pdf_file_path}")

    except FileNotFoundError:
        print(f'The file {pdf_file_path} does not exist.')
    except Exception as e:
        print(f'An unexpected error with the file {pdf_file_path} occurred: {str(e)}')


def rename_files() -> None:
    """
    Assures file names consistency by renaming certain files
    """

    for file in os.listdir(PDFs_DIR_PATH):
        filepath = os.path.join(PDFs_DIR_PATH, file)

        if os.path.isfile(filepath):
            # Replace whitespaces with underscores in the filename
            s = file.replace(' ', '-')
            new_filename = s
            if s.count('.') > 1:
                s_without_periods = s.replace('.', '')
                last_period_position = s.rfind('.')
                new_filename = s_without_periods[:last_period_position] + '.' + s_without_periods[last_period_position:]

            if new_filename.endswith('p.df'):
                new_filename = new_filename[:-4] + '.pdf'
            new_filepath = os.path.join(PDFs_DIR_PATH, new_filename)

            os.rename(filepath, new_filepath)
            print(f'Renamed file {filepath} - {new_filepath}')

    print('Finished renaming files')


def find_pds_with_multiple_invoices(directory_path: str) -> list[str]:
    """
    Searches for PDF files within the specified directory that contain multiple occurrences of 
    certain keywords related to invoices. These keywords are 'commercial invoice', 'factura comercial',
    'payment conditions', 'condiciones de pago', and 'FOB'. A PDF is considered a match if it contains 
    at least two occurrences of each of these terms.

    Parameters:
    directory_path (str): The file path of the directory.

    Returns:
    list[str]: A list of file paths for the PDFs.
    """

    matching_pdfs = []

    for filename in os.listdir(directory_path):
        filepath = os.path.join(directory_path, filename)

        if os.path.isfile(filepath) and filename.lower().endswith('.pdf'):
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                occurrences_commercialinvoice = 0
                occurrences_paymentconditions = 0
                occurrences_fob = 0

                for page_num in range(len(pdf_reader.pages)):
                    text = pdf_reader.pages[page_num].extract_text()
                    text = text.replace(' ', '').lower()
                    occurrences_commercialinvoice += text.count('commercialinvoice') + text.count('facturacomercial')
                    occurrences_paymentconditions += text.count('paymentconditions') + text.count('condicionesdepago')
                    occurrences_fob += text.count('fob')

                if occurrences_commercialinvoice >= 2 and occurrences_paymentconditions >= 2 and occurrences_fob >= 2:
                    matching_pdfs.append(filepath)

    return matching_pdfs


def get_cell_content(docx_path: str, table_index: int, row_index: int, col_index: int) -> str:
    """
    Gets the content of a specific cell from a table in a DOCX document. The table, row, 
    and column are specified by their indices.

    Parameters:
    docx_path (str): The file path of the DOCX document.
    table_index (int): The index of the table within the document.
    row_index (int): The row index within the table.
    col_index (int): The column index within the row.

    Returns:
    str: The text content of the specified cell. 
    """

    doc = Document(docx_path)

    if table_index >= len(doc.tables):
        return "Table index out of range!"

    table = doc.tables[table_index]

    if row_index >= len(table.rows) or col_index >= len(table.columns):
        return "Row or Column index out of range!"

    cell = table.cell(row_index, col_index)
    return cell.text.strip()


def find_cell_with_exact_content(docx_path: str, target_string: str) -> list[int] or None:
    """
    Searches through (case-insensitive) all tables in a DOCX document to find a cell that exactly matches the given target string.

    Parameters:
    docx_path (str): The file path of the DOCX document.
    target_string (str): The string to search for.

    Returns:
    list[int] or None: A list containing the indices of the table, row, and column of the matching cell, if found.
    Returns None if no matching cell is found.
    """

    doc = Document(docx_path)

    for table_num, table in enumerate(doc.tables):
        for row_num, row in enumerate(table.rows):
            for col_num, cell in enumerate(row.cells):
                cell_content = cell.text.strip()  
                if cell_content.lower() == target_string.lower():
                    return [table_num, row_num, col_num]

    return None


def merge_row_cells_with_below(docx_path: str, table_index: int, row_index: int, save_path: str = None) -> str or None:
    """
    Merges each cell in a specified row of a table with the cell directly below it in a DOCX document. 

    Parameters:
    docx_path (str): The file path of the DOCX document.
    table_index (int): The index of the table within the document.
    row_index (int): The index of the row within the table to merge with the row below.
    save_path (str): The file path to save the modified document. If None, the document is not saved.

    Returns:
    str or None: Returns an error message if the table or row indices are out of range, or None if the operation is successful.
    """

    doc = Document(docx_path)

    if table_index >= len(doc.tables):
        return "Table index out of range!"

    table = doc.tables[table_index]

    if row_index >= len(table.rows) - 1:
        return "Row index out of range!"

    for col_index in range(len(table.columns)):
        cell = table.cell(row_index, col_index)
        cell_below = table.cell(row_index + 1, col_index)

        try:
            cell.merge(cell_below)
        except:
            continue

    if save_path:
        doc.save(save_path)
        print(f'Cells merged in {save_path}')









